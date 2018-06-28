#! /usr/bin/env python3

from docx import Document
from docx.shared import Inches
import os

recordset = [
    [1, 2, 3],
    ['hello', 'world', 'how are you?'],
    [1230, 'mixed listing', 'here we go'],
]

doc_name = 'demo.docx'

if os.path.exists(doc_name):
    document = Document(doc_name)
else:
    document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('images/monty-truth.jpg', width=Inches(1.25))

document.add_heading('Table Examples', level=1)

table_types = [
  'Colorful Grid',
  'Colorful List',
  'Colorful Shading',
  'Dark List',
  'Light Grid',
  'Light List',
  'Light Shading',
  'Medium Grid 1',
  'Medium Grid 2',
  'Medium Grid 3',
  'Medium List 1',
  'Medium List 2',
  'Medium Shading 1',
  'Medium Shading 2',
]

# Create an example of all the tables
for table_type in table_types:
    for idx in range(1, 6):
        style = '%s - %s' % (table_type, 'Accent %d' % (idx))
        document.add_heading(style, level=2)
        try:
            table = document.add_table(rows=1, cols=3, style=style.replace(' ', ''))
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Qty'
            hdr_cells[1].text = 'Id'
            hdr_cells[2].text = 'Desc'
            for item in recordset:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item[0])
                row_cells[1].text = str(item[1])
                row_cells[2].text = str(item[2])
        except:
            error = 'Failed to process: %s' % (style)
            document.add_paragraph(error)
            print(error)

document.add_page_break()

document.save(doc_name)
